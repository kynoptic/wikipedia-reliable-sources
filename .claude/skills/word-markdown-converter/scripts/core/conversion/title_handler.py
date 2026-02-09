#!/usr/bin/env python3
"""
Title Detection and Processing for Word-to-Markdown Conversion.

This module provides intelligent title detection and handling for Word documents,
including smart title extraction from content and document properties.

Key Components:
    - TitleHandler: Main class for title detection and processing
    - Smart title detection from first substantial paragraph
    - Fallback to document properties when no content title found
    - Support for various Word title styles and formats

Features:
    - Explicit Title style detection
    - Content-based title candidate identification
    - Document property title extraction
    - Title length normalization and validation
    - Robust error handling for corrupted documents

Dependencies:
    - python-docx: Word document processing
    - logging: Error and debug logging
"""

import logging
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


class TitleHandler:
    """Handles title detection and processing for Word documents."""

    def __init__(self, logger: logging.Logger):
        """Initialize the title handler.

        Args:
            logger: Logger instance for error and debug messages
        """
        self.logger = logger

    def detect_and_process_title(
        self, word_document: Document
    ) -> tuple[Optional[str], Optional[Paragraph]]:
        """Detect and process document title from Word document.

        Returns a tuple of (title_markdown, skip_paragraph) where:
        - title_markdown: The markdown title string (e.g., "# Title") or None
        - skip_paragraph: The paragraph to skip during content processing or None

        Args:
            word_document: The Word document to process

        Returns:
            Tuple of (title_markdown, paragraph_to_skip)
        """
        has_explicit_title = self._has_explicit_title_style(word_document)

        if has_explicit_title:
            return None, None

        # Try content-based title detection first
        title_candidate = self._find_title_candidate(word_document)
        if title_candidate:
            title_text = title_candidate.text.strip()
            return f"# {title_text}", title_candidate

        # Check if document starts with Heading 1 - if so, don't use property title
        if self._starts_with_heading_1(word_document):
            return None, None

        # Fallback to document properties
        property_title = self._extract_document_property_title(word_document)
        if property_title and property_title != "Untitled Document":
            return f"# {property_title}", None

        return None, None

    def _has_explicit_title_style(self, word_document: Document) -> bool:
        """Check if document has explicit Title style paragraphs.

        Args:
            word_document: The Word document to check

        Returns:
            True if document has explicit Title style with content
        """
        try:
            return any(
                getattr(paragraph.style, "name", "") == "Title"
                and paragraph.text
                and paragraph.text.strip()
                for paragraph in word_document.paragraphs
            )
        except Exception as e:
            self.logger.debug(f"Error checking explicit title: {e}")
            return False

    def _find_title_candidate(self, word_document: Document) -> Optional[Paragraph]:
        """Find first substantial paragraph that could serve as document title.

        Strategy: Find the first substantial paragraph that isn't already a
        heading/list to handle cases where Word documents lack explicit Title
        styles but have clear title content as the first meaningful paragraph.

        Args:
            word_document: The Word document to analyze

        Returns:
            Paragraph that could serve as title, or None if not found
        """
        for paragraph in word_document.paragraphs:
            style_name = self._get_safe_style_name(paragraph)
            paragraph_text = paragraph.text or ""

            # Skip empty paragraphs - they can't be meaningful titles
            if not paragraph_text.strip():
                continue

            # Stop scanning at first heading/subtitle to avoid using body content as title
            if self._is_heading_style(style_name):
                return None

            # Skip list paragraphs as they're unlikely to be document titles
            if self._is_list_style(style_name):
                continue

            # Found our title candidate - first substantial non-heading paragraph
            return paragraph

        return None

    def _get_safe_style_name(self, paragraph: Paragraph) -> str:
        """Safely extract style name from paragraph.

        Args:
            paragraph: The paragraph to extract style from

        Returns:
            Style name or "Normal" if extraction fails
        """
        try:
            return paragraph.style.name if paragraph.style else "Normal"
        except Exception as e:
            self.logger.debug(f"Error extracting style name: {e}")
            return "Normal"

    def _starts_with_heading_1(self, word_document: Document) -> bool:
        """Check if document starts with a Heading 1 paragraph.

        Args:
            word_document: The Word document to check

        Returns:
            True if first non-empty paragraph is Heading 1
        """
        for paragraph in word_document.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text or not paragraph.text.strip():
                continue
            # Check if first non-empty paragraph is Heading 1
            style_name = self._get_safe_style_name(paragraph)
            return style_name == "Heading 1"
        return False

    def _is_heading_style(self, style_name: str) -> bool:
        """Check if style name represents a heading style.

        Args:
            style_name: The style name to check

        Returns:
            True if style represents a heading
        """
        return style_name.startswith("Heading") or style_name == "Subtitle" or style_name == "Title"

    def _is_list_style(self, style_name: str) -> bool:
        """Check if style name represents a list style.

        Args:
            style_name: The style name to check

        Returns:
            True if style represents a list
        """
        return "List" in style_name

    def _extract_document_property_title(self, doc: Document) -> str:
        """Extract document title from properties or content fallback.

        Args:
            doc: The Word document to extract title from

        Returns:
            Document title or "Untitled Document" if none found
        """
        try:
            # Try document properties first
            if hasattr(doc.core_properties, "title") and doc.core_properties.title:
                return doc.core_properties.title

            # Try to get title from first heading
            for para in doc.paragraphs:
                if para.style and para.style.name in ["Title", "Heading 1"] and para.text.strip():
                    return para.text.strip()

            # Fallback to first non-empty paragraph
            for para in doc.paragraphs:
                if para.text.strip():
                    title = para.text.strip()
                    # Truncate long titles
                    return title[:100] + "..." if len(title) > 100 else title

            return "Untitled Document"

        except Exception as e:
            self.logger.debug(f"Error extracting document title: {e}")
            return "Untitled Document"
