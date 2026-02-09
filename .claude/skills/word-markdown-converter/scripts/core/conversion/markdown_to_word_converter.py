#!/usr/bin/env python3
"""
Markdown-to-Word Document Converter with Style Mapping.

This module provides comprehensive conversion of Markdown documents (.md)
to Microsoft Word format (.docx) with sophisticated style mapping and
document structure preservation. It supports configurable style mappings,
intelligent content classification, and high-fidelity conversion workflows.

Key Components:
    - MarkdownToWordConverter: Main conversion engine with style awareness
    - Style mapping from Markdown elements to Word styles
    - Configurable style mapping system with YAML configuration
    - Image embedding and reference handling
    - Table conversion with proper Word formatting
    - Hyperlink and cross-reference preservation

Conversion Features:
    - Document structure preservation (headings, lists, tables)
    - Code block conversion with syntax highlighting support
    - Image embedding with proper sizing and positioning
    - YAML frontmatter extraction and metadata handling
    - List formatting with proper indentation
    - Table handling with Word-native formatting

Configuration:
    - Style mappings defined in tools/style_map.yml
    - Customizable paragraph and character styles
    - Template document support for consistent formatting
    - Image resolution and sizing preferences

Dependencies:
    - python-docx: Word document creation and style management
    - markdown: Markdown parsing and AST generation
    - PyYAML: Configuration file parsing and processing
    - pandoc: Advanced markdown parsing and conversion
    - pathlib: Modern path handling utilities

Usage:
    converter = MarkdownToWordConverter(logger)
    success = converter.convert_document('input.md', 'output.docx')
"""

import logging
import re
import subprocess
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Pt, RGBColor


class MarkdownToWordConverter:
    """Convert Markdown documents to Word format with style preservation."""

    def __init__(self, logger: Optional[logging.Logger] = None):
        """Initialize converter with optional logger."""
        self.logger = logger or logging.getLogger(__name__)
        # Support both bundled skill location and original repo location
        skill_config = Path(__file__).parent.parent.parent.parent / "config" / "style_map.yml"
        repo_config = Path(__file__).parent.parent.parent / "tools" / "style_map.yml"
        self.style_map_path = skill_config if skill_config.exists() else repo_config
        self.style_map = self._load_style_map()
        self.reverse_style_map = self._create_reverse_style_map()

        # Default reference template with HMS IT theme
        self.default_template = (
            Path(__file__).parent.parent.parent / "tools" / "reference_template.docx"
        )

    def _load_style_map(self) -> Dict[str, Any]:
        """Load style mapping configuration from YAML file."""
        try:
            with open(self.style_map_path, "r") as f:
                return yaml.safe_load(f)
        except Exception as e:
            self.logger.error(f"Failed to load style map: {e}")
            return {"character_styles": {}, "paragraph_styles": {}}

    def _create_reverse_style_map(self) -> Dict[str, str]:
        """Create reverse mapping from Markdown to Word styles."""
        reverse_map = {}

        # Map paragraph styles
        for word_style, config in self.style_map.get("paragraph_styles", {}).items():
            md_type = config.get("markdown", "")
            if md_type.startswith("h") and md_type[1:].isdigit():
                reverse_map[md_type] = word_style
            elif md_type == "italic":
                reverse_map["subtitle"] = word_style

        # Map character styles
        for word_style, config in self.style_map.get("character_styles", {}).items():
            md_type = config.get("markdown", "")
            if md_type == "backticks":
                reverse_map["code"] = word_style
            elif md_type == "kbd_tag":
                reverse_map["kbd"] = word_style

        return reverse_map

    def convert_document(
        self, input_path: str, output_path: str, template_path: Optional[str] = None
    ) -> bool:
        """
        Convert a Markdown file to Word document.

        Args:
            input_path: Path to input Markdown file
            output_path: Path to output Word document
            template_path: Optional template document for styling.
                          If None, uses default HMS IT theme template.

        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            input_file = Path(input_path)
            output_file = Path(output_path)

            if not input_file.exists():
                self.logger.error(f"Input file not found: {input_path}")
                return False

            # Create output directory if needed
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Use original file directly - no preprocessing needed
            # Pandoc handles blockquotes as indented paragraphs naturally
            temp_input = input_file

            # Determine which template to use
            if template_path:
                # User-specified template
                template_to_use = template_path
            elif self.default_template.exists():
                # Use default HMS IT theme template
                template_to_use = str(self.default_template)
                self.logger.debug("Using default HMS IT theme template")
            else:
                # No template available
                template_to_use = None

            # Create Word document
            if template_to_use and Path(template_to_use).exists():
                doc = Document(template_to_use)
                self.logger.info(f"Using template: {Path(template_to_use).name}")
            else:
                doc = Document()
                self._setup_default_styles(doc)

            # Convert using Pandoc for better fidelity
            success = self._convert_with_pandoc(temp_input, output_file, template_to_use)

            if success:
                # Post-process with python-docx for style mapping
                self._apply_custom_styles(output_file)
                self.logger.info(f"Successfully converted {input_path} to {output_path}")
                return True
            else:
                # Fallback to manual conversion
                with open(input_file, "r", encoding="utf-8") as f:
                    md_content = f.read()
                self._convert_manually(md_content, doc, input_file.parent)
                doc.save(output_file)
                self.logger.info(f"Converted {input_path} to {output_path} (fallback method)")
                return True

        except Exception as e:
            self.logger.error(f"Conversion failed: {e}", exc_info=True)
            return False

    def _convert_with_pandoc(
        self, input_file: Path, output_file: Path, template_path: Optional[str] = None
    ) -> bool:
        """Use Pandoc for high-fidelity conversion."""
        try:
            # Build markdown format string with required extensions
            md_extensions = [
                "yaml_metadata_block",
                "pipe_tables",
                "fenced_code_blocks",
                "raw_html",
                "lists_without_preceding_blankline",
            ]
            md_format = "markdown+" + "+".join(md_extensions)

            cmd = [
                "pandoc",
                str(input_file),
                "-o",
                str(output_file),
                "-f",
                md_format,
                "-t",
                "docx",
                "--wrap=preserve",
                "--no-highlight",  # Disable syntax highlighting in code blocks
            ]

            # Add Lua filter for heading style mapping (first h1 -> Title, etc.)
            heading_filter_path = (
                Path(__file__).parent.parent / "filters" / "md_to_word_headings.lua"
            )
            if heading_filter_path.exists():
                cmd.extend(["--lua-filter", str(heading_filter_path)])
                self.logger.debug(f"Using heading style filter: {heading_filter_path}")

            # Add Lua filter for <kbd> tag handling
            filter_path = Path(__file__).parent.parent / "filters" / "kbd_style.lua"
            if filter_path.exists():
                cmd.extend(["--lua-filter", str(filter_path)])
                self.logger.debug(f"Using kbd style filter: {filter_path}")

            # Add Lua filter to fix blockquotes nested in lists without blank lines
            # Makes conversion robust for real-world markdown files
            blockquote_fix_filter = (
                Path(__file__).parent.parent / "filters" / "fix_list_blockquotes.lua"
            )
            if blockquote_fix_filter.exists():
                cmd.extend(["--lua-filter", str(blockquote_fix_filter)])
                self.logger.debug(f"Using blockquote fix filter: {blockquote_fix_filter}")

            # Note: Blockquotes are handled naturally by Pandoc as indented paragraphs
            # No special processing needed - they will be indented by default

            if template_path:
                cmd.extend(["--reference-doc", template_path])

            # Add resource path for images
            resource_path = input_file.parent
            cmd.extend(["--resource-path", str(resource_path)])

            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode != 0:
                self.logger.warning(f"Pandoc conversion warning: {result.stderr}")
                return False

            return True

        except FileNotFoundError:
            self.logger.warning("Pandoc not found, using fallback conversion")
            return False
        except Exception as e:
            self.logger.error(f"Pandoc conversion error: {e}")
            return False

    def _convert_manually(self, md_content: str, doc: Document, resource_path: Path):
        """Manual conversion when Pandoc is not available.

        Implements the same heading style mapping as the Pandoc Lua filter:
        - All h1 before first h2 → Title style
        - h1 after first h2 → Heading 1 style
        - h2 → Heading 1 style (offset from Title)
        - h3 → Heading 2 style (and so on)
        """
        lines = md_content.split("\n")
        in_code_block = False
        seen_h2_or_below = False

        for line in lines:
            # Skip empty lines in code blocks
            if in_code_block and not line.strip():
                continue

            # Code blocks
            if line.startswith("```"):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                self._add_code_paragraph(doc, line)
                continue

            # Headings with Title style mapping
            if line.startswith("#"):
                seen_h2_or_below = self._process_heading(doc, line, seen_h2_or_below)
                continue

            # Lists
            if self._is_list_line(line):
                self._add_list_item(doc, line)
                continue

            # Tables (basic support)
            if self._is_table_line(line):
                continue

            # Regular paragraphs
            if line.strip():
                p = doc.add_paragraph(line)
                self._apply_inline_formatting(p, line)

    def _add_code_paragraph(self, doc: Document, line: str):
        """Add a code paragraph to the document."""
        p = doc.add_paragraph(line)
        p.style = "Code"

    def _process_heading(self, doc: Document, line: str, seen_h2_or_below: bool) -> bool:
        """Process a heading line and apply appropriate style.

        Returns:
            bool: Updated seen_h2_or_below flag
        """
        level = len(line) - len(line.lstrip("#"))
        heading_text = line[level:].strip()

        # Track when we see first h2 or below
        if level >= 2 and not seen_h2_or_below:
            seen_h2_or_below = True

        # Apply heading style mapping
        if level == 1 and not seen_h2_or_below:
            # h1 before first h2 → Title style
            p = doc.add_paragraph(heading_text)
            p.style = "Title"
        elif level == 1:
            # h1 after first h2 → Heading 1
            doc.add_heading(heading_text, level=1)
        elif level >= 2:
            # h2+ → offset by one level (h2→Heading 1, h3→Heading 2, etc.)
            doc.add_heading(heading_text, level=level - 1)

        return seen_h2_or_below

    def _is_list_line(self, line: str) -> bool:
        """Check if line is a list item."""
        return bool(re.match(r"^(\s*)[*+-]\s+", line) or re.match(r"^(\s*)\d+\.\s+", line))

    def _add_list_item(self, doc: Document, line: str):
        """Add a list item to the document."""
        list_text = re.sub(r"^(\s*)[*+-]\s+", "", line)
        list_text = re.sub(r"^(\s*)\d+\.\s+", "", list_text)
        doc.add_paragraph(list_text, style="List Bullet")

    def _is_table_line(self, line: str) -> bool:
        """Check if line is part of a table."""
        return "|" in line and line.strip().startswith("|")

    def _apply_inline_formatting(self, paragraph, text: str):
        """Apply inline formatting like bold, italic, code."""
        # Clear existing runs by removing them from the paragraph's XML
        for run in paragraph.runs[:]:
            run._element.getparent().remove(run._element)

        # Pattern for inline code
        code_pattern = r"`([^`]+)`"
        # Pattern for bold
        bold_pattern = r"\*\*([^*]+)\*\*"
        # Pattern for italic
        italic_pattern = r"\*([^*]+)\*"

        # Process text for formatting
        parts = []
        last_end = 0

        # Find all formatting matches
        for match in re.finditer(f"({code_pattern}|{bold_pattern}|{italic_pattern})", text):
            # Add text before match
            if match.start() > last_end:
                parts.append(("normal", text[last_end : match.start()]))

            # Determine formatting type
            if match.group().startswith("`"):
                parts.append(("code", match.group(2)))
            elif match.group().startswith("**"):
                parts.append(("bold", match.group(3)))
            elif match.group().startswith("*"):
                parts.append(("italic", match.group(4)))

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            parts.append(("normal", text[last_end:]))

        # Apply formatting to runs
        for fmt_type, content in parts:
            run = paragraph.add_run(content)
            if fmt_type == "code":
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif fmt_type == "bold":
                run.bold = True
            elif fmt_type == "italic":
                run.italic = True

    def _setup_default_styles(self, doc: Document):
        """Setup default styles for the document."""
        # Ensure basic styles exist
        styles = doc.styles

        # Code character style
        if "Code" not in [s.name for s in styles]:
            code_style = styles.add_style("Code", WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = "Courier New"
            code_style.font.size = Pt(10)

        # Input character style for keyboard input
        if "Input" not in [s.name for s in styles]:
            input_style = styles.add_style("Input", WD_STYLE_TYPE.CHARACTER)
            input_style.font.name = "Courier New"
            input_style.font.size = Pt(10)
            input_style.font.bold = True

        # List styles
        if "List Bullet" not in [s.name for s in styles]:
            try:
                list_style = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
                list_style.paragraph_format.left_indent = Inches(0.25)
            except Exception:
                pass  # Style might already exist

    def _apply_custom_styles(self, docx_path: Path):
        """Apply custom style mappings to existing document.

        Post-processes Pandoc output to apply character styles from style_map.yml:
        - Replaces Pandoc's "Verbatim Char" with "Code" character style
        - Replaces Pandoc's "SourceCode" with template's "Source Code" paragraph style
        - Detects keyboard input markers and applies "Input" character style
        """
        try:
            self.logger.debug(f"Post-processing styles in {docx_path}")
            doc = Document(docx_path)

            # Ensure our custom character styles exist
            self._ensure_character_styles_exist(doc)

            # Apply character and paragraph styles
            code_count = 0
            kbd_count = 0
            source_code_count = 0

            for paragraph in doc.paragraphs:
                # Remap Pandoc's "SourceCode" paragraph style to template's "Source Code"
                if paragraph.style.name == "SourceCode":
                    try:
                        paragraph.style = "Source Code"
                        source_code_count += 1
                    except KeyError:
                        self.logger.warning("'Source Code' style not found in template")

                # First pass: detect and style keyboard input markers
                self._apply_keyboard_input_styles(paragraph)

                # Second pass: apply code styles
                for run in paragraph.runs:
                    # Skip empty runs
                    if not run.text or not run.text.strip():
                        continue

                    # Replace Pandoc's "Verbatim Char" with our "Code" style
                    if self._is_pandoc_code_style(run):
                        self.logger.debug(f"Found Verbatim Char run: '{run.text[:30]}'")
                        run.style = "Code"
                        code_count += 1
                        self.logger.debug(f"Applied 'Code' style to: {run.text[:30]}")

            # Apply blockquote indentation
            blockquote_count = self._apply_blockquote_indentation(doc)

            # Save modified document
            doc.save(docx_path)
            if code_count > 0 or kbd_count > 0 or source_code_count > 0 or blockquote_count > 0:
                self.logger.info(
                    f"Applied styles to {docx_path.name}: "
                    f"{code_count} inline code, {kbd_count} keyboard input, "
                    f"{source_code_count} code blocks, {blockquote_count} blockquotes indented"
                )

        except Exception as e:
            self.logger.warning(f"Could not apply custom styles: {e}")

    def _apply_keyboard_input_styles(self, paragraph):
        """Detect and style keyboard input markers from Lua filter.

        Finds runs containing ⌘⌘text⌘⌘ markers and splits them to apply
        the "Input" character style to the keyboard input text.
        """
        import re

        # Pattern to match our keyboard input markers
        kbd_pattern = r"⌘⌘([^⌘]+)⌘⌘"

        runs_to_process = list(paragraph.runs)
        for run in runs_to_process:
            if not run.text or "⌘⌘" not in run.text:
                continue

            # Find all keyboard input markers in this run
            matches = list(re.finditer(kbd_pattern, run.text))
            if not matches:
                continue

            # Split the run and apply styles
            text = run.text
            parts = []
            last_end = 0

            for match in matches:
                # Add text before marker
                if match.start() > last_end:
                    parts.append(("normal", text[last_end : match.start()]))

                # Add keyboard input text
                parts.append(("input", match.group(1)))
                last_end = match.end()

            # Add remaining text
            if last_end < len(text):
                parts.append(("normal", text[last_end:]))

            # Clear the run and rebuild with proper styles
            run.text = ""
            for part_type, part_text in parts:
                new_run = paragraph.add_run(part_text)

                # Copy formatting from original run
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size

                # Apply keyboard input style
                if part_type == "input":
                    new_run.style = "Input"
                    self.logger.debug(f"Applied 'Input' style to: {part_text}")

    def _ensure_character_styles_exist(self, doc: Document):
        """Ensure required character styles exist in document."""
        styles = doc.styles

        # Code character style
        if "Code" not in [s.name for s in styles]:
            code_style = styles.add_style("Code", WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = "Courier New"
            code_style.font.size = Pt(10)
            self.logger.debug("Created 'Code' character style")

        # Input character style
        if "Input" not in [s.name for s in styles]:
            input_style = styles.add_style("Input", WD_STYLE_TYPE.CHARACTER)
            input_style.font.name = "Courier New"
            input_style.font.size = Pt(10)
            input_style.font.bold = True
            self.logger.debug("Created 'Input' character style")

        # Hyperlink character style
        if "Hyperlink" not in [s.name for s in styles]:
            hyperlink_style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
            hyperlink_style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)  # Blue color
            hyperlink_style.font.underline = WD_UNDERLINE.SINGLE
            self.logger.debug("Created 'Hyperlink' character style")

    def _is_pandoc_code_style(self, run) -> bool:
        """Check if run uses Pandoc's default code style.

        Pandoc converts inline code (backticks) to runs with "Verbatim Char" style.
        Detect these to replace with our "Code" character style from style_map.yml.
        """
        try:
            if not run.style:
                return False

            style_name = run.style.name
            # Pandoc uses "Verbatim Char" for inline code
            return style_name == "Verbatim Char"

        except Exception:
            return False

    def _apply_blockquote_indentation(self, doc: Document) -> int:
        """Apply indentation to blockquote paragraphs and nested lists.

        Pandoc converts BlockQuote elements to paragraphs with "Blockquote" or "Block Text" style.
        This method adds left indentation and removes list numbering to make blockquotes
        visually distinct. Also indents list items immediately following blockquotes
        (which are lists nested within the blockquote).

        Returns:
            int: Number of paragraphs indented
        """
        count = 0
        last_was_blockquote = False

        for i, paragraph in enumerate(doc.paragraphs):
            # Pandoc uses "Blockquote" or "Block Text" styles for actual blockquotes
            # (Markdown content starting with >)
            style_name = paragraph.style.name

            if style_name in ["Blockquote", "Block Text"]:
                # Note: "Body Text" is NOT included because Pandoc uses it for regular
                # continuation paragraphs, not blockquotes. Only "Blockquote" and "Block Text"
                # indicate actual blockquote content (Markdown lines starting with >).

                # Remove numbering properties if present (blockquotes shouldn't be list items)
                num_pr = paragraph._element.xpath(".//w:numPr")
                if num_pr:
                    for num_element in num_pr:
                        num_element.getparent().remove(num_element)
                    self.logger.debug(
                        f"Removed numbering from blockquote: {paragraph.text[:50]}..."
                    )

                # Change style to Normal (blockquotes use indentation only, not special styles)
                paragraph.style = "Normal"

                # Apply left indentation (0.5 inches = 36 points)
                paragraph.paragraph_format.left_indent = Pt(36)
                count += 1
                last_was_blockquote = True
                self.logger.debug(f"Applied blockquote indentation to: {paragraph.text[:50]}...")

            elif last_was_blockquote:
                # Check if this is a list item (has numbering properties)
                has_numbering = bool(paragraph._element.xpath(".//w:numPr"))

                if has_numbering:
                    # This is a list item inside the blockquote
                    # Indent it to align with or slightly beyond the blockquote
                    # Blockquote has 36pt base, list items get 72pt (36pt base + 36pt list indent)
                    current_indent = paragraph.paragraph_format.left_indent
                    if current_indent is None or current_indent.pt < 60:
                        paragraph.paragraph_format.left_indent = Pt(72)
                        count += 1
                        self.logger.debug(
                            f"Applied blockquote list indent to: {paragraph.text[:50]}..."
                        )
                else:
                    # Not a list item, blockquote context ends
                    last_was_blockquote = False
            else:
                last_was_blockquote = False

        return count

    def batch_convert(
        self, input_dir: str, output_dir: str, template_path: Optional[str] = None
    ) -> Tuple[int, int]:
        """
        Batch convert all Markdown files in a directory.

        Args:
            input_dir: Directory containing Markdown files
            output_dir: Directory for output Word documents
            template_path: Optional template document

        Returns:
            Tuple of (successful_conversions, failed_conversions)
        """
        input_path = Path(input_dir)
        output_path = Path(output_dir)

        if not input_path.exists():
            self.logger.error(f"Input directory not found: {input_dir}")
            return (0, 0)

        # Create output directory
        output_path.mkdir(parents=True, exist_ok=True)

        # Find all Markdown files
        md_files = list(input_path.glob("**/*.md"))

        successful = 0
        failed = 0

        for md_file in md_files:
            # Calculate relative path for output
            rel_path = md_file.relative_to(input_path)
            output_file = output_path / rel_path.with_suffix(".docx")

            # Create output subdirectory if needed
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Convert file
            if self.convert_document(str(md_file), str(output_file), template_path):
                successful += 1
                self.logger.info(f"Converted: {md_file.name}")
            else:
                failed += 1
                self.logger.error(f"Failed: {md_file.name}")

        self.logger.info(f"Batch conversion complete: {successful} successful, {failed} failed")
        return (successful, failed)


if __name__ == "__main__":
    # Test conversion
    logging.basicConfig(level=logging.INFO)
    converter = MarkdownToWordConverter()

    # Example single file conversion
    # converter.convert_document("input.md", "output.docx")

    # Example batch conversion
    # converter.batch_convert("content/markdown", "content/docx_output")
