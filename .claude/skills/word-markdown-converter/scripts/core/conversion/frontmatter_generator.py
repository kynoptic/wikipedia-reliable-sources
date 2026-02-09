#!/usr/bin/env python3
"""
Frontmatter Generation for Word-to-Markdown Conversion.

This module provides YAML frontmatter generation for converted Markdown documents,
including metadata extraction, source tracking, and conversion statistics.

Key Components:
    - FrontmatterGenerator: Main class for YAML frontmatter generation
    - Source document metadata and hash calculation
    - Conversion statistics collection and formatting
    - YAML frontmatter structure with round-trip validation

Features:
    - SHA-256 hash calculation for source documents
    - Relative path handling for portability
    - Comprehensive conversion statistics
    - Document metadata extraction from Word properties
    - YAML frontmatter with proper formatting
    - Round-trip validation support

Dependencies:
    - python-docx: Word document processing
    - PyYAML: YAML generation and formatting
    - pathlib: Modern path handling
    - hashlib: File hash calculation
    - datetime: Timestamp generation
"""

import hashlib
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

import yaml
from docx import Document


class FrontmatterGenerator:
    """Handles YAML frontmatter generation for converted Markdown documents."""

    def __init__(self, logger: logging.Logger):
        """Initialize the frontmatter generator.

        Args:
            logger: Logger instance for error and debug messages
        """
        self.logger = logger

    def generate_frontmatter(
        self,
        docx_path: str,
        output_path: str,
        doc: Document,
        style_stats: Dict[str, int],
        image_mapping: Dict[str, str],
    ) -> str:
        """Generate YAML frontmatter with source path, hash, and conversion stats.

        Args:
            docx_path: Path to the source Word document
            output_path: Path to the output Markdown file
            doc: The Word document object
            style_stats: Style statistics from conversion
            image_mapping: Image mapping dictionary

        Returns:
            YAML frontmatter string with metadata
        """
        try:
            # Calculate source document hash
            source_hash = self._calculate_file_hash(docx_path)

            # Get relative paths for better portability
            try:
                source_relative = Path(docx_path).relative_to(Path.cwd())
            except ValueError:
                source_relative = Path(docx_path)

            try:
                output_relative = Path(output_path).relative_to(Path.cwd())
            except ValueError:
                output_relative = Path(output_path)

            # Collect enhanced conversion statistics
            conversion_stats = self._collect_conversion_stats(doc, style_stats)

            # Build frontmatter data
            frontmatter_data = {
                "source": {
                    "path": str(source_relative),
                    "hash": source_hash,
                    "last_modified": self._get_file_timestamp(docx_path),
                },
                "conversion": {
                    "timestamp": datetime.now().isoformat(),
                    "converter_version": "1.0.0",
                    "output_path": str(output_relative),
                },
                "document": {
                    "title": self._extract_document_title(doc),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "images": len(image_mapping),
                },
                "styles": conversion_stats,
                "round_trip": {"source_hash": source_hash, "expected_match": True},
            }

            # Convert to YAML frontmatter format
            yaml_content = yaml.dump(frontmatter_data, default_flow_style=False, sort_keys=False)

            return f"---\n{yaml_content}---"

        except Exception as e:
            self.logger.warning(f"Error generating frontmatter: {e}")
            # Return minimal frontmatter
            return (
                f"---\nsource: {docx_path}\nconversion_timestamp: {datetime.now().isoformat()}\n---"
            )

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of the source file.

        Args:
            file_path: Path to the file to hash

        Returns:
            Hexadecimal hash string
        """
        try:
            with open(file_path, "rb") as f:
                file_content = f.read()
                return hashlib.sha256(file_content).hexdigest()
        except Exception as e:
            self.logger.warning(f"Could not calculate file hash: {e}")
            return "unknown"

    def _get_file_timestamp(self, file_path: str) -> str:
        """Get last modified timestamp of file.

        Args:
            file_path: Path to the file

        Returns:
            ISO format timestamp string
        """
        try:
            mtime = os.path.getmtime(file_path)
            return datetime.fromtimestamp(mtime).isoformat()
        except Exception as e:
            self.logger.warning(f"Could not get file timestamp: {e}")
            return datetime.now().isoformat()

    def _collect_conversion_stats(
        self, doc: Document, style_stats: Dict[str, int]
    ) -> Dict[str, Any]:
        """Collect detailed conversion statistics.

        Args:
            doc: The Word document
            style_stats: Style statistics from conversion

        Returns:
            Dictionary with conversion statistics
        """
        stats = {
            "character_styles": {},
            "paragraph_styles": {},
            "total_runs": 0,
            "hyperlinks_found": 0,
            "code_blocks_found": 0,
            "lists_found": 0,
        }

        try:
            # Count character styles
            char_styles = {
                k.replace("char_style_", ""): v
                for k, v in style_stats.items()
                if k.startswith("char_style_")
            }
            stats["character_styles"] = char_styles

            # Count paragraph styles
            para_styles = {
                k.replace("para_style_", ""): v
                for k, v in style_stats.items()
                if k.startswith("para_style_")
            }
            stats["paragraph_styles"] = para_styles

            # Count specific elements
            for para in doc.paragraphs:
                # Count runs
                stats["total_runs"] += len(para.runs)

                # Count hyperlinks
                hyperlinks = para._element.xpath(".//w:hyperlink")
                stats["hyperlinks_found"] += len(hyperlinks)

                # Count code blocks
                if para.style and "code" in para.style.name.lower():
                    stats["code_blocks_found"] += 1

                # Count list items
                if para.style and "list" in para.style.name.lower():
                    stats["lists_found"] += 1

        except Exception as e:
            self.logger.debug(f"Error collecting conversion stats: {e}")

        return stats

    def _extract_document_title(self, doc: Document) -> str:
        """Extract document title from properties or content fallback.

        Args:
            doc: The Word document

        Returns:
            Document title or "Untitled Document" if none found
        """
        try:
            # Try document properties first
            if hasattr(doc.core_properties, "title") and doc.core_properties.title:
                return doc.core_properties.title

            # Try to get title from first heading
            for para in doc.paragraphs:
                if para.style and para.style.name in ["Title", "Heading 1"] and para.text.strip():
                    return para.text.strip()

            # Fallback to first non-empty paragraph
            for para in doc.paragraphs:
                if para.text.strip():
                    title = para.text.strip()
                    # Truncate long titles
                    return title[:100] + "..." if len(title) > 100 else title

            return "Untitled Document"

        except Exception as e:
            self.logger.debug(f"Error extracting document title: {e}")
            return "Untitled Document"
