#!/usr/bin/env python3
"""
Create reference.docx with custom paragraph and character styles
Used by Pandoc for MD→DOCX conversion with proper style mapping
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def create_custom_reference_docx():
    """Create reference.docx with custom styles"""

    doc = Document()
    styles = doc.styles

    _create_character_styles(styles)
    _create_paragraph_styles(styles)
    _add_sample_content(doc, styles)

    output_path = Path(__file__).parent / "reference.docx"
    doc.save(str(output_path))
    print(f"\nReference document saved to: {output_path}")

    return str(output_path)


def _create_character_styles(styles):
    """Create custom character styles for inline formatting"""
    _create_code_character_style(styles)
    _create_input_character_style(styles)


def _create_code_character_style(styles):
    """Create Code character style for inline code"""
    try:
        code_style = styles.add_style("Code", WD_STYLE_TYPE.CHARACTER)
        code_font = code_style.font
        code_font.name = "Aptos Mono"
        code_font.size = Pt(10.5)

        _add_background_shading(code_style._element, "F2F2F2")

        print("Created 'Code' character style (Aptos Mono 10.5pt, #F2F2F2 background)")
    except ValueError:
        print("'Code' style already exists, skipping")


def _create_input_character_style(styles):
    """Create Input character style for keyboard input"""
    try:
        input_style = styles.add_style("Input", WD_STYLE_TYPE.CHARACTER)
        input_font = input_style.font
        input_font.name = "Aptos Mono"
        input_font.size = Pt(10.0)

        _add_background_shading(input_style._element, "F2F2F2")

        print("Created 'Input' character style (Aptos Mono 10.0pt, #F2F2F2 background)")
    except ValueError:
        print("'Input' style already exists, skipping")


def _add_background_shading(style_element, fill_color):
    """Add background shading to a style element"""
    rPr = style_element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_element.append(rPr)

    shd = OxmlElement("w:shd")
    shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", fill_color)
    shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "clear")
    rPr.append(shd)


def _create_paragraph_styles(styles):
    """Create custom paragraph styles"""
    _create_code_block_style(styles)
    _create_list_styles(styles)
    _create_block_text_style(styles)
    _create_subtitle_style(styles)
    _create_toc_styles(styles)
    _create_heading_styles(styles)


def _create_code_block_style(styles):
    """Create Code block paragraph style"""
    try:
        code_block_style = styles.add_style("Code block", WD_STYLE_TYPE.PARAGRAPH)
        code_block_font = code_block_style.font
        code_block_font.name = "Consolas"
        code_block_font.size = Pt(10)
        code_block_para = code_block_style.paragraph_format
        code_block_para.left_indent = Pt(36)
        code_block_para.space_before = Pt(6)
        code_block_para.space_after = Pt(6)
        print("Created 'Code block' paragraph style")
    except ValueError:
        print("'Code block' style already exists, skipping")


def _create_list_styles(styles):
    """Create list-related paragraph styles"""
    list_styles = [
        ("List Paragraph", {"left_indent": Pt(36)}),
        ("List Bullet", {"left_indent": Pt(36), "first_line_indent": Pt(-18)}),
    ]

    for style_name, para_format in list_styles:
        try:
            list_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            list_style.font.name = "Calibri"
            list_style.font.size = Pt(11)
            for attr, value in para_format.items():
                setattr(list_style.paragraph_format, attr, value)
            print(f"Created '{style_name}' paragraph style")
        except ValueError:
            print(f"'{style_name}' style already exists, skipping")


def _create_block_text_style(styles):
    """Create Block Text paragraph style"""
    try:
        block_text_style = styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
        block_text_font = block_text_style.font
        block_text_font.name = "Calibri"
        block_text_font.size = Pt(11)
        block_text_para = block_text_style.paragraph_format
        block_text_para.left_indent = Pt(36)
        block_text_para.right_indent = Pt(36)
        block_text_para.space_before = Pt(6)
        block_text_para.space_after = Pt(6)
        print("Created 'Block Text' paragraph style")
    except ValueError:
        print("'Block Text' style already exists, skipping")


def _create_subtitle_style(styles):
    """Create Subtitle paragraph style"""
    try:
        subtitle_style = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = "Calibri"
        subtitle_font.size = Pt(14)
        subtitle_font.italic = True
        subtitle_font.color.rgb = RGBColor(68, 68, 68)
        print("Created 'Subtitle' paragraph style")
    except ValueError:
        print("'Subtitle' style already exists, skipping")


def _create_toc_styles(styles):
    """Create TOC paragraph styles"""
    for i in range(1, 3):
        toc_name = f"toc {i}"
        try:
            toc_style = styles.add_style(toc_name, WD_STYLE_TYPE.PARAGRAPH)
            toc_font = toc_style.font
            toc_font.name = "Calibri"
            toc_font.size = Pt(11)
            if i == 1:
                toc_font.bold = True
            toc_para = toc_style.paragraph_format
            toc_para.left_indent = Pt(18 * i)
            print(f"Created '{toc_name}' paragraph style")
        except ValueError:
            print(f"'{toc_name}' style already exists, skipping")


def _create_heading_styles(styles):
    """Create heading paragraph styles (H1-H6)"""
    for i in range(1, 7):
        heading_name = f"Heading {i}"
        try:
            heading_style = styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = "Calibri"
            heading_font.size = Pt(20 - (i * 2))
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(43, 87, 154)
            heading_para = heading_style.paragraph_format
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
            print(f"Created '{heading_name}' paragraph style")
        except ValueError:
            print(f"'{heading_name}' style already exists, skipping")


def _add_sample_content(doc, styles):
    """Add sample content to demonstrate styles"""
    doc.add_heading("Custom Reference Document", 0)

    intro = doc.add_paragraph()
    intro.add_run("This is the reference document for Word-to-Markdown conversions. ")
    intro.add_run("It defines the paragraph and character styles used by Pandoc when converting ")
    intro.add_run("Markdown back to Word format.")

    for i in range(1, 7):
        doc.add_heading(f"Heading Level {i} Sample", i)
        doc.add_paragraph(f"This is sample content under heading level {i}.")

    doc.add_heading("Character Style Examples", 1)

    code_para = doc.add_paragraph("Inline code example: ")
    code_run = code_para.add_run('<div class="example">')
    code_run.style = styles["Code"]
    code_para.add_run(" should appear with Code character style.")

    input_para = doc.add_paragraph("Keyboard input example: Press ")
    input_run = input_para.add_run("Ctrl+S")
    input_run.style = styles["Input"]
    input_para.add_run(" to save the document.")

    doc.add_heading("Code Block Example", 1)

    doc.add_paragraph("function example() {", style="Code block")
    doc.add_paragraph('  return "Hello, World!";', style="Code block")
    doc.add_paragraph("}", style="Code block")


if __name__ == "__main__":
    create_custom_reference_docx()
